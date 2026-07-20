from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "BATTLE_FES_開会式進行台本.docx"

# compact_reference_guide preset + named brand override
FONT_JP = "Yu Gothic"
NAVY = "0B2545"
BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
GOLD = "C68B00"
ORANGE = "D95E1E"
RED = "9B1C1C"
INK = "1E2329"
MUTED = "66717D"
LIGHT_BLUE = "E8EEF5"
LIGHT_GOLD = "FFF6D9"
LIGHT_ORANGE = "FCECE5"
LIGHT_GRAY = "F4F6F9"
WHITE = "FFFFFF"
TABLE_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def set_style_font(style, name=FONT_JP, size=None, color=None, bold=None):
    style.font.name = name
    fonts = style._element.get_or_add_rPr().get_or_add_rFonts()
    fonts.set(qn("w:ascii"), name)
    fonts.set(qn("w:hAnsi"), name)
    fonts.set(qn("w:eastAsia"), name)
    if size is not None:
        style.font.size = Pt(size)
    if color is not None:
        style.font.color.rgb = rgb(color)
    if bold is not None:
        style.font.bold = bold


def set_run_font(run, name=FONT_JP, size=None, color=None, bold=None, italic=None):
    run.font.name = name
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), name)
    rfonts.set(qn("w:hAnsi"), name)
    rfonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_fill(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for edge, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        tag = qn(f"w:{edge}")
        node = tc_mar.find(tag)
        if node is None:
            node = OxmlElement(f"w:{edge}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_borders(table, color="D7DEE7", size=6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(size))
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def set_table_geometry(table, widths_dxa, indent_dxa=TABLE_INDENT_DXA):
    assert sum(widths_dxa) == TABLE_WIDTH_DXA
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr

    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(TABLE_WIDTH_DXA))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            width = widths_dxa[index]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(width))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(width / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_row_cant_split(row):
    tr_pr = row._tr.get_or_add_trPr()
    cant_split = OxmlElement("w:cantSplit")
    cant_split.set(qn("w:val"), "true")
    tr_pr.append(cant_split)


def set_keep_with_next(paragraph, value=True):
    paragraph.paragraph_format.keep_with_next = value


def add_page_number(paragraph):
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), "PAGE")
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "1"
    run.append(text)
    field.append(run)
    paragraph._p.append(field)


def add_header_footer(section):
    header_p = section.header.paragraphs[0]
    header_p.paragraph_format.space_after = Pt(0)
    header_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = header_p.add_run("BATTLE FES 2026")
    set_run_font(left, size=8.5, color=MUTED, bold=True)
    right = header_p.add_run("\t開会式進行台本")
    set_run_font(right, size=8.5, color=MUTED)

    footer_p = section.footer.paragraphs[0]
    footer_p.paragraph_format.space_before = Pt(0)
    footer_p.paragraph_format.tab_stops.add_tab_stop(Inches(6.5), WD_TAB_ALIGNMENT.RIGHT)
    left = footer_p.add_run("司会読み上げ用")
    set_run_font(left, size=8.5, color=MUTED)
    page = footer_p.add_run("\tPAGE ")
    set_run_font(page, size=8.5, color=MUTED)
    add_page_number(footer_p)


def configure_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.72)
    section.right_margin = Inches(0.8)
    section.bottom_margin = Inches(0.72)
    section.left_margin = Inches(0.8)
    section.header_distance = Inches(0.36)
    section.footer_distance = Inches(0.36)
    add_header_footer(section)

    normal = doc.styles["Normal"]
    set_style_font(normal, size=11, color=INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    title = doc.styles["Title"]
    set_style_font(title, size=26, color=NAVY, bold=True)
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(6)
    title.paragraph_format.line_spacing = 1.0

    subtitle = doc.styles["Subtitle"]
    set_style_font(subtitle, size=13.5, color=MUTED)
    subtitle.paragraph_format.space_before = Pt(0)
    subtitle.paragraph_format.space_after = Pt(14)
    subtitle.paragraph_format.line_spacing = 1.0

    h1 = doc.styles["Heading 1"]
    set_style_font(h1, size=16, color=BLUE, bold=True)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    set_style_font(h2, size=13, color=ORANGE, bold=True)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.keep_with_next = True

    h3 = doc.styles["Heading 3"]
    set_style_font(h3, size=12, color=DARK_BLUE, bold=True)
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(5)
    h3.paragraph_format.keep_with_next = True

    script = doc.styles.add_style("Script Block", 1)
    set_style_font(script, size=11.4, color=INK)
    script.paragraph_format.left_indent = Inches(0.14)
    script.paragraph_format.right_indent = Inches(0.05)
    script.paragraph_format.space_before = Pt(0)
    script.paragraph_format.space_after = Pt(10)
    script.paragraph_format.line_spacing = 1.35
    script.paragraph_format.keep_together = True

    kicker = doc.styles.add_style("Kicker", 1)
    set_style_font(kicker, size=9, color=GOLD, bold=True)
    kicker.paragraph_format.space_before = Pt(0)
    kicker.paragraph_format.space_after = Pt(4)

    settings = doc.settings._element
    theme_font = settings.find(qn("w:themeFontLang"))
    if theme_font is None:
        theme_font = OxmlElement("w:themeFontLang")
        settings.append(theme_font)
    theme_font.set(qn("w:val"), "ja-JP")
    theme_font.set(qn("w:eastAsia"), "ja-JP")


def set_paragraph_shading(paragraph, fill=LIGHT_GOLD, border=GOLD):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)

    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    left = borders.find(qn("w:left"))
    if left is None:
        left = OxmlElement("w:left")
        borders.append(left)
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), "18")
    left.set(qn("w:space"), "8")
    left.set(qn("w:color"), border)


def add_callout(doc, label, text, fill=LIGHT_GOLD, border=GOLD, label_color=GOLD):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    set_table_borders(table, border, 8)
    cell = table.cell(0, 0)
    set_cell_fill(cell, fill)
    set_cell_margins(cell, top=90, start=180, bottom=90, end=180)
    set_row_cant_split(table.rows[0])
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    label_run = p.add_run(label)
    set_run_font(label_run, size=8.5, color=label_color, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    text_run = p2.add_run(text)
    set_run_font(text_run, size=10.8, color=INK, bold=True)
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(0)
    after.paragraph_format.space_before = Pt(0)
    after.paragraph_format.line_spacing = Pt(3)
    set_run_font(after.add_run(" "), size=2, color=WHITE)
    return table


def add_metric_strip(doc):
    table = doc.add_table(rows=1, cols=4)
    set_table_geometry(table, [2340, 2340, 2340, 2340])
    set_table_borders(table, NAVY, 6)
    metrics = [
        ("19:45", "開会式開始"),
        ("20:00", "リレー開始"),
        ("20:45", "投票受付"),
        ("22:15", "本投票"),
    ]
    for cell, (value, label) in zip(table.rows[0].cells, metrics):
        set_cell_fill(cell, NAVY)
        set_cell_margins(cell, top=70, start=110, bottom=70, end=110)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(value)
        set_run_font(run, size=14, color="FFD65A", bold=True)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(0)
        run2 = p2.add_run(label)
        set_run_font(run2, size=8.5, color=WHITE, bold=True)
    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_after = Pt(0)
    spacer.paragraph_format.line_spacing = Pt(4)
    set_run_font(spacer.add_run(" "), size=2, color=WHITE)


def add_run_of_show(doc):
    heading = doc.add_heading("15分の進行", level=1)
    heading.paragraph_format.space_before = Pt(9)
    heading.paragraph_format.space_after = Pt(5)
    table = doc.add_table(rows=1, cols=3)
    set_table_geometry(table, [950, 1350, 7060])
    set_table_borders(table)
    headers = ["経過", "時刻", "内容"]
    for cell, text in zip(table.rows[0].cells, headers):
        set_cell_fill(cell, NAVY)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text != "内容" else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=9.5, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])

    rows = [
        ("0:00", "19:45", "オープニング・開催挨拶"),
        ("1:00", "19:46", "イベント概要"),
        ("2:30", "19:47:30", "3チーム・出演者紹介"),
        ("5:00", "19:50", "チーム投票の説明"),
        ("8:00", "19:53", "個人賞・貫通BONUSの説明"),
        ("10:30", "19:55:30", "リレー配信の移動方法"),
        ("12:30", "19:57:30", "最終案内・最初の出演者紹介"),
        ("14:00", "19:59", "次枠への移動時間"),
    ]
    for index, row_values in enumerate(rows):
        row = table.add_row()
        set_row_cant_split(row)
        for column, (cell, value) in enumerate(zip(row.cells, row_values)):
            if index % 2 == 1:
                set_cell_fill(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if column < 2 else WD_ALIGN_PARAGRAPH.LEFT
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(value)
            set_run_font(run, size=9.6, color=INK, bold=(column < 2))
    set_table_geometry(table, [950, 1350, 7060])
    for row_index, row in enumerate(table.rows):
        vertical_margin = 42 if row_index else 54
        for cell in row.cells:
            set_cell_margins(cell, top=vertical_margin, start=120, bottom=vertical_margin, end=120)
    return table


def add_script_section(doc, number, time_range, heading, lines, fill=LIGHT_GOLD, border=GOLD):
    h = doc.add_heading(f"{number}  |  {time_range}  |  {heading}", level=2)
    set_keep_with_next(h)
    p = doc.add_paragraph(style="Script Block")
    set_paragraph_shading(p, fill=fill, border=border)
    for index, line in enumerate(lines):
        if index:
            p.add_run().add_break()
        run = p.add_run(line)
        set_run_font(run, size=11.4, color=INK, bold=False)
    return p


def add_checklist_table(doc):
    heading = doc.add_heading("開会式直前チェック", level=1)
    heading.paragraph_format.page_break_before = True
    table = doc.add_table(rows=1, cols=2)
    set_table_geometry(table, [1100, 8260])
    set_table_borders(table)
    for cell, text in zip(table.rows[0].cells, ("確認", "項目")):
        set_cell_fill(cell, LIGHT_BLUE)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER if text == "確認" else WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(text)
        set_run_font(run, size=10, color=NAVY, bold=True)
    set_repeat_table_header(table.rows[0])
    set_row_cant_split(table.rows[0])

    items = [
        "19:45に開会式を開始する",
        "3チーム・9名を紹介する",
        "投票受付は20:45〜22:30と案内する",
        "本投票は22:15〜22:30と案内する",
        "優勝チームへの投票は1人1票と案内する",
        "3つの個人賞を案内する",
        "ルーム背景の文字と貫通BONUSを説明する",
        "フォロー一覧の一番上から次枠へ移動と案内する",
        "20:00の最初の出演者は犬飼音子（ねこ）さんと紹介する",
        "19:59までに話し終える",
    ]
    for index, item in enumerate(items):
        row = table.add_row()
        set_row_cant_split(row)
        if index % 2 == 1:
            for cell in row.cells:
                set_cell_fill(cell, LIGHT_GRAY)
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        box = p0.add_run("□")
        set_run_font(box, size=13, color=GOLD, bold=True)
        p1 = row.cells[1].paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        run = p1.add_run(item)
        set_run_font(run, size=10.2, color=INK)
    set_table_geometry(table, [1100, 8260])
    return table


def audit_document(doc):
    section = doc.sections[0]
    assert round(section.page_width.inches, 3) == 8.5
    assert round(section.page_height.inches, 3) == 11.0
    assert round(section.top_margin.inches, 3) == 0.72
    assert round(section.right_margin.inches, 3) == 0.8
    assert round(section.bottom_margin.inches, 3) == 0.72
    assert round(section.left_margin.inches, 3) == 0.8
    assert round(section.header_distance.inches, 3) == 0.36
    assert round(section.footer_distance.inches, 3) == 0.36

    for table in doc.tables:
        tbl_pr = table._tbl.tblPr
        tbl_w = tbl_pr.find(qn("w:tblW"))
        tbl_ind = tbl_pr.find(qn("w:tblInd"))
        assert tbl_w is not None and tbl_w.get(qn("w:w")) == str(TABLE_WIDTH_DXA)
        assert tbl_ind is not None and tbl_ind.get(qn("w:w")) == str(TABLE_INDENT_DXA)
        grid_widths = [int(col.get(qn("w:w"))) for col in table._tbl.tblGrid]
        assert sum(grid_widths) == TABLE_WIDTH_DXA
        for row in table.rows:
            for index, cell in enumerate(row.cells):
                tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
                assert tc_w is not None and int(tc_w.get(qn("w:w"))) == grid_widths[index]


def build():
    doc = Document()
    configure_document(doc)

    kicker = doc.add_paragraph(style="Kicker")
    kicker.add_run("RUN OF SHOW  |  2026.07.18")
    title = doc.add_paragraph("BATTLE FES 2026\n開会式進行台本", style="Title")
    title.paragraph_format.keep_with_next = True
    subtitle = doc.add_paragraph("19:45〜20:00  /  司会読み上げ用", style="Subtitle")
    subtitle.paragraph_format.keep_with_next = True
    add_metric_strip(doc)
    add_callout(
        doc,
        "最重要",
        "14分以内に案内を終える。最後の1分は、20:00の次枠への移動・トラブル対応用に空ける。",
        fill=LIGHT_GOLD,
        border=GOLD,
        label_color=GOLD,
    )
    add_run_of_show(doc)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(7)
    note.paragraph_format.space_after = Pt(0)
    run = note.add_run("次枠：20:00  TEAM NOVA  犬飼音子（ねこ）さん")
    set_run_font(run, size=11, color=RED, bold=True)

    front_heading = doc.add_heading("読み上げ台本  |  前半", level=1)
    front_heading.paragraph_format.page_break_before = True
    add_script_section(
        doc,
        "1",
        "0:00〜1:00",
        "オープニング",
        [
            "皆さん、こんばんは。",
            "ColorSing Summer Event 2026「BATTLE FES」へようこそ！",
            "本日は、CRIMSON、NOVA、GOLDENの3チーム、総勢9名によるチーム対抗歌唱バトルをお届けします。",
        ],
    )
    add_script_section(
        doc,
        "2",
        "1:00〜2:30",
        "イベント概要",
        [
            "これから9名の出演者が、1人15分ずつリレー形式で配信します。",
            "最終的に、リスナー投票、個人賞加点、そして各枠で獲得したライブスコアを合計し、優勝チームを決定します。",
        ],
        fill=LIGHT_BLUE,
        border=BLUE,
    )
    add_script_section(
        doc,
        "3",
        "2:30〜5:00",
        "チーム・出演者紹介",
        [
            "TEAM CRIMSONは、んごご、PK、リーダーのまぐろふぉん。",
            "TEAM NOVAは、犬飼音子（ねこ）、なぽる、リーダーのりんか。",
            "TEAM GOLDENは、潮てら、あわ、リーダーのiran。",
            "コメント欄で、応援しているチームの名前をぜひ教えてください！",
            "それでは、チームコールをお願いします。CRIMSON！ NOVA！ GOLDEN！",
        ],
        fill=LIGHT_GRAY,
        border=DARK_BLUE,
    )
    add_script_section(
        doc,
        "4",
        "5:00〜8:00",
        "チーム投票の説明",
        [
            "優勝チームへの投票は、1人1票です。",
            "投票受付は20時45分から22時30分まで、BATTLE FES特設サイトで行います。",
            "投票は時間が進むほど1票のポイントが大きくなり、22時15分から始まる本投票では、最大の5000ポイントになります。",
            "本投票は22時30分までの15分間です。",
        ],
        fill=LIGHT_ORANGE,
        border=ORANGE,
    )

    back_heading = doc.add_heading("読み上げ台本  |  後半", level=1)
    back_heading.paragraph_format.page_break_before = True
    add_script_section(
        doc,
        "5",
        "8:00〜9:00",
        "個人賞の説明",
        [
            "チーム投票とは別に、MVP（最優秀歌唱賞）、最優秀エンタメ賞、ベストモーメント賞の3つの個人賞があります。",
            "ぜひ9名全員の配信を見て、それぞれの賞にふさわしい出演者を選んでください。",
        ],
        fill=LIGHT_BLUE,
        border=BLUE,
    )
    add_script_section(
        doc,
        "6",
        "9:00〜10:30",
        "貫通BONUSの説明",
        [
            "今回は「貫通BONUS」があります。",
            "各出演者のルーム背景に文字が表示されます。9つの枠を回って文字を集め、キーワードを完成させてください。",
            "完成したキーワードを本投票中に入力して正解すると、チーム投票に5000ポイントのボーナスが加算されます。",
            "キーワードを見逃さないよう、各出演者のルーム背景にも注目してください。",
        ],
        fill=LIGHT_GOLD,
        border=GOLD,
    )
    add_script_section(
        doc,
        "7",
        "10:30〜12:30",
        "リレー配信の移動方法",
        [
            "各出演者の持ち時間は15分です。",
            "配信が終わったら、その出演者のフォロー一覧を開いてください。",
            "フォロー一覧の一番上に次の出演者が表示されています。そこから次の枠へ移動できます。",
            "配信中にも定期コメントでご案内します。",
        ],
        fill=LIGHT_GRAY,
        border=DARK_BLUE,
    )
    add_script_section(
        doc,
        "8",
        "12:30〜14:00",
        "最初の出演者へ",
        [
            "それでは、20時からBATTLE FESのリレー配信が始まります。",
            "最初の出演者は、TEAM NOVAの犬飼音子（ねこ）さんです。",
            "歌、コメント、スーパーいいね、そして皆さんの1票で、最後まで一緒にBATTLE FESを盛り上げてください。",
            "それでは、開戦です！",
            "次の配信は、フォロー一覧の一番上から移動してください。",
        ],
        fill=LIGHT_ORANGE,
        border=ORANGE,
    )

    add_checklist_table(doc)
    doc.add_heading("開会式では言わないこと", level=1)
    add_callout(
        doc,
        "非公開・説明不要",
        "個人賞の加点ポイント数 / 不正対策・重複判定の詳しい仕組み / 完成したキーワードの答え",
        fill="FBEAEA",
        border=RED,
        label_color=RED,
    )
    closing = doc.add_paragraph()
    closing.paragraph_format.space_before = Pt(16)
    closing.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = closing.add_run("19:59までに話し終え、20:00の次枠へつなぐ")
    set_run_font(run, size=13, color=NAVY, bold=True)

    doc.core_properties.title = "BATTLE FES 2026 開会式進行台本"
    doc.core_properties.subject = "2026年7月18日 19:45〜20:00 司会読み上げ用"
    doc.core_properties.author = "BATTLE FES 2026 運営"
    doc.core_properties.keywords = "BATTLE FES, 開会式, 進行台本"

    audit_document(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"saved {OUTPUT}")


if __name__ == "__main__":
    build()
